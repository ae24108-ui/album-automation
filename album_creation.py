import os
from docx import Document
from docx.shared import Mm
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.table import WD_ROW_HEIGHT_RULE

#設定(入力フォルダ、出力フォルダのパスを入力してください)
IMAGE_FOLDER = r""
OUTPUT_FILE = r""

ROWS = 3
COLS = 3
CELL_WIDTH_MM = 55
CELL_HEIGHT_MM = 75

def create_album():
    #画像パスをimagesに入れる
    images = []
    file_list = os.listdir(IMAGE_FOLDER) #listdirectory
    file_list = sorted(file_list)

    for f in file_list:
        if f.lower().endswith((".jpg", ".jpeg", ".png")):
                full_path = os.path.join(IMAGE_FOLDER, f)
                images.append(full_path)

    if not images:
        print("No images found")
        return

    doc = Document()

    #ページの余白を'狭い'に設定
    margin = Mm(12.7)

    section = doc.sections[0]
    section.top_margin = margin
    section.bottom_margin = margin
    section.left_margin = margin
    section.right_margin = margin

    image_index = 0 #imagesの中身の番号(最初は0番目)
    total_images = len(images)

    while image_index < total_images:
        #表作成
        table = doc.add_table(rows=ROWS, cols=COLS)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER #表を中央揃え
        table.autofit = False #表の自動サイズ調整をOFF

        #セルサイズ固定
        for row in table.rows:
            row.height = Mm(CELL_HEIGHT_MM)
            row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY #高さを固定値化

        for col in table.columns:
            col.width = Mm(CELL_WIDTH_MM)
            for cell in col.cells:
                cell.width = Mm(CELL_WIDTH_MM)

        #写真挿入(cellを指定→paragraphを作成→runを作成し,pictureを挿入)
        for row in range(ROWS):
            for col in range(COLS):
                if image_index >= total_images:
                    break

                cell = table.cell(row,col)
                paragraph = cell.paragraphs[0]#セル内の一行目
                run = paragraph.add_run()

                run.add_picture(
                    images[image_index], #imagesの0個目を入れる
                    width=Mm(CELL_WIDTH_MM)
                )

                image_index += 1

        if image_index < total_images:
            doc.add_page_break()

    doc.save(OUTPUT_FILE)
    print('Done')

if __name__ == '__main__':
    create_album()